"""
LifeOps API — Word Document (.docx) Report Generator
=====================================================
Generates executive-level Word reports with custom styles,
data tables, KPI blocks, and professional formatting.
"""
import io
import datetime as dt
from typing import Dict, Any, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


# ── Color Palette ─────────────────────────────────
COLOR_PRIMARY_HEX = "0B0F17"    # Dark obsidian
COLOR_EMERALD_HEX = "10B981"    # Emerald green
COLOR_CYAN_HEX = "06B6D4"       # Cyan
COLOR_PURPLE_HEX = "8B5CF6"     # Purple
COLOR_BG_LIGHT_HEX = "F8FAFC"   # Light slate
COLOR_BORDER_HEX = "E2E8F0"     # Border gray
COLOR_MUTED_HEX = "64748B"      # Text muted

RGB_PRIMARY = RGBColor(11, 15, 23)
RGB_EMERALD = RGBColor(16, 185, 129)
RGB_CYAN = RGBColor(6, 182, 212)
RGB_PURPLE = RGBColor(139, 92, 246)
RGB_MUTED = RGBColor(100, 116, 139)
RGB_DARK = RGBColor(30, 41, 59)


# ── XML Helper Functions for Table Cell Styling ───

def set_cell_background(cell, fill_hex: str):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner padding of a table cell (in twips, 20 twips = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_heading_with_accent(doc: Document, text: str, level=1):
    """Add a clean styled heading with a bottom border or accent color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True

    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.bold = True

    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGB_PRIMARY
    elif level == 2:
        run.font.size = Pt(12.5)
        run.font.color.rgb = RGB_DARK
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGB_MUTED

    return p


# ── Report Templates Implementation ───────────────

def generate_docx_report(
    template_type: str,
    user_email: str,
    date_from: dt.date,
    date_to: dt.date,
    data: Dict[str, Any],
) -> io.BytesIO:
    """
    Main generator dispatching to specific template builders.
    Returns in-memory BytesIO stream containing the .docx file.
    """
    doc = Document()

    # Page Margins (2 cm on all sides)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    if template_type == "sport_performance":
        _build_sport_report(doc, user_email, date_from, date_to, data)
    elif template_type == "project_status":
        _build_project_report(doc, user_email, date_from, date_to, data)
    else:
        # Default: Monthly / Integral Summary
        _build_monthly_integral_report(doc, user_email, date_from, date_to, data)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _build_header(doc: Document, title: str, subtitle: str, user_email: str, date_from: dt.date, date_to: dt.date):
    """Build standardized executive header banner."""
    # Header Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_title = title_p.add_run(f"⚡ LifeOps — {title}")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGB_PRIMARY

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run(subtitle)
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(10.5)
    run_sub.font.color.rgb = RGB_MUTED

    # Meta Info Table (User, Period, Date)
    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_data = [
        ("👤 USUARIO", user_email),
        ("📅 PERIODO", f"{date_from.strftime('%d/%m/%Y')} - {date_to.strftime('%d/%m/%Y')}"),
        ("⏱️ GENERADO", dt.datetime.now().strftime('%d/%m/%Y %H:%M')),
    ]

    widths = [Inches(2.3), Inches(2.5), Inches(2.2)]
    for i, (label, val) in enumerate(meta_data):
        cell = meta_table.cell(0, i)
        cell.width = widths[i]
        set_cell_background(cell, COLOR_BG_LIGHT_HEX)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r_lbl = p.add_run(f"{label}\n")
        r_lbl.font.size = Pt(7.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGB_MUTED

        r_val = p.add_run(val)
        r_val.font.size = Pt(9.5)
        r_val.font.bold = True
        r_val.font.color.rgb = RGB_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── Template 1: Monthly / Integral Summary ────────

def _build_monthly_integral_report(doc: Document, user_email: str, date_from: dt.date, date_to: dt.date, data: Dict[str, Any]):
    _build_header(
        doc,
        title="Informe Mensual Integral",
        subtitle="Resumen 360° de rendimiento personal, deportivo, lecturas y avance profesional.",
        user_email=user_email,
        date_from=date_from,
        date_to=date_to,
    )

    activities = data.get("activities", [])
    workouts = [a for a in activities if a.get("activity_type") == "sport"]
    books = [a for a in activities if a.get("activity_type") == "book"]
    films = [a for a in activities if a.get("activity_type") == "film"]
    tasks = data.get("tasks", [])
    projects = data.get("projects", [])

    # ── KPI Summary Cards ──
    add_heading_with_accent(doc, "1. Resumen Ejecutivo (KPIs del Periodo)", level=1)

    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_table.autofit = False

    total_km = sum(w.get("workout", {}).get("distance_km") or 0 for w in workouts)
    total_sport_min = sum(w.get("duration_minutes") or 0 for w in workouts)
    completed_books = sum(1 for b in books if b.get("book", {}).get("status") == "completed")
    done_tasks = sum(1 for t in tasks if t.get("status") == "done")

    kpis = [
        ("🏃 Total Deporte", f"{total_km:.1f} km", f"{len(workouts)} sesiones"),
        ("⏱️ Tiempo Activo", f"{total_sport_min / 60:.1f} hrs", f"{total_sport_min} min totales"),
        ("📚 Libros Leídos", f"{completed_books}", f"{len(books)} en biblioteca"),
        ("✅ Tareas Hechas", f"{done_tasks}", f"{len(tasks)} totales"),
    ]

    for i, (label, val, sub) in enumerate(kpis):
        cell = kpi_table.cell(0, i)
        cell.width = Inches(1.75)
        set_cell_background(cell, COLOR_BG_LIGHT_HEX)
        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r1 = p.add_run(f"{label}\n")
        r1.font.size = Pt(8)
        r1.font.color.rgb = RGB_MUTED

        r2 = p.add_run(f"{val}\n")
        r2.font.size = Pt(14)
        r2.font.bold = True
        r2.font.color.rgb = RGB_PRIMARY

        r3 = p.add_run(sub)
        r3.font.size = Pt(7.5)
        r3.font.color.rgb = RGB_MUTED

    # ── Section: Deporte ──
    add_heading_with_accent(doc, "2. Deporte & Actividad Física", level=1)
    if workouts:
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Fecha", "Entrenamiento", "Tipo", "Distancia / Duración", "Calorías / PB"]
        hdr_row = table.rows[0]
        for idx, text in enumerate(headers):
            cell = hdr_row.cells[idx]
            set_cell_background(cell, COLOR_PRIMARY_HEX)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

        for act in workouts:
            w = act.get("workout", {})
            row = table.add_row()
            cells = row.cells
            
            dist_str = f"{w.get('distance_km')} km" if w.get('distance_km') else f"{act.get('duration_minutes', 0)} min"
            pb_str = "🏅 Récord PB" if w.get('personal_best') else (f"{w.get('calories')} kcal" if w.get('calories') else "-")

            vals = [
                act.get("date", "-"),
                act.get("title", "-"),
                w.get("workout_type", "sport").capitalize(),
                dist_str,
                pb_str,
            ]
            for idx, text in enumerate(vals):
                cell = cells[idx]
                set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(str(text))
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGB_DARK
    else:
        p = doc.add_paragraph("No se registraron entrenamientos en este periodo.")
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGB_MUTED

    # ── Section: Biblioteca & Cine ──
    add_heading_with_accent(doc, "3. Lectura, Cine & Cultura", level=1)
    if books or films:
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Categoría", "Título", "Detalle / Autor", "Estado / Valoración"]
        for idx, text in enumerate(headers):
            cell = table.rows[0].cells[idx]
            set_cell_background(cell, COLOR_PRIMARY_HEX)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

        for b in books:
            bk = b.get("book", {})
            row = table.add_row()
            rating_str = f"★ {b.get('rating')}/5" if b.get('rating') else bk.get("status", "reading")
            vals = ["📖 Libro", b.get("title"), bk.get("author", "Autor no especificado"), rating_str]
            for idx, text in enumerate(vals):
                cell = row.cells[idx]
                set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(str(text))
                r.font.size = Pt(8.5)

        for f in films:
            fl = f.get("film", {})
            row = table.add_row()
            rating_str = f"★ {f.get('rating')}/5" if f.get('rating') else "-"
            vals = [f"🎬 {fl.get('media_type', 'Film').capitalize()}", f.get("title"), fl.get("platform", "Cine"), rating_str]
            for idx, text in enumerate(vals):
                cell = row.cells[idx]
                set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(str(text))
                r.font.size = Pt(8.5)
    else:
        p = doc.add_paragraph("Sin registros de lectura o cine en este periodo.")
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGB_MUTED

    # ── Section: Profesional & Proyectos ──
    add_heading_with_accent(doc, "4. Portafolio de Proyectos & Tareas", level=1)
    if tasks:
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Tarea", "Prioridad", "Fecha Límite", "Estado"]
        for idx, text in enumerate(headers):
            cell = table.rows[0].cells[idx]
            set_cell_background(cell, COLOR_PRIMARY_HEX)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

        for t in tasks:
            row = table.add_row()
            status_trans = {
                "todo": "📝 Por Hacer",
                "in_progress": "⚡ En Curso",
                "review": "🔍 En Revisión",
                "done": "✅ Completada"
            }.get(t.get("status"), t.get("status"))

            vals = [
                t.get("title"),
                (t.get("priority") or "medium").capitalize(),
                t.get("due_date") or "-",
                status_trans,
            ]
            for idx, text in enumerate(vals):
                cell = row.cells[idx]
                set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(str(text))
                r.font.size = Pt(8.5)
    else:
        p = doc.add_paragraph("No hay tareas registradas para el periodo.")
        p.runs[0].font.italic = True


# ── Template 2: Sport Report ───────────────────────

def _build_sport_report(doc: Document, user_email: str, date_from: dt.date, date_to: dt.date, data: Dict[str, Any]):
    _build_header(
        doc,
        title="Dossier de Rendimiento Deportivo",
        subtitle="Métricas avanzadas de entrenamientos, volumen por disciplina y récords personales.",
        user_email=user_email,
        date_from=date_from,
        date_to=date_to,
    )
    workouts = [a for a in data.get("activities", []) if a.get("activity_type") == "sport"]

    add_heading_with_accent(doc, "1. Desglose de Entrenamientos", level=1)
    if workouts:
        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Fecha", "Título", "Deporte", "Distancia", "Duración", "Calorías"]
        for idx, text in enumerate(headers):
            cell = table.rows[0].cells[idx]
            set_cell_background(cell, COLOR_PRIMARY_HEX)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

        for act in workouts:
            w = act.get("workout", {})
            row = table.add_row()
            vals = [
                act.get("date", "-"),
                act.get("title", "-"),
                w.get("workout_type", "sport").capitalize(),
                f"{w.get('distance_km')} km" if w.get("distance_km") else "-",
                f"{act.get('duration_minutes', 0)} min",
                f"{w.get('calories')} kcal" if w.get("calories") else "-",
            ]
            for idx, text in enumerate(vals):
                cell = row.cells[idx]
                set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(str(text))
                r.font.size = Pt(8.5)


# ── Template 3: Project Status Report ──────────────

def _build_project_report(doc: Document, user_email: str, date_from: dt.date, date_to: dt.date, data: Dict[str, Any]):
    _build_header(
        doc,
        title="Informe de Estado de Proyectos & Tareas",
        subtitle="Seguimiento de portafolio, control de horas y estado de entregables.",
        user_email=user_email,
        date_from=date_from,
        date_to=date_to,
    )
    projects = data.get("projects", [])
    tasks = data.get("tasks", [])

    add_heading_with_accent(doc, "1. Portafolio de Proyectos Activos", level=1)
    if projects:
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Proyecto", "Estado", "Presupuesto", "Fecha Objetivo"]
        for idx, text in enumerate(headers):
            cell = table.rows[0].cells[idx]
            set_cell_background(cell, COLOR_PRIMARY_HEX)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(255, 255, 255)

        for prj in projects:
            row = table.add_row()
            budget_str = f"{prj.get('budget'):,} €".replace(",", ".") if prj.get("budget") else "Sin definir"
            vals = [
                prj.get("name"),
                prj.get("status", "active").capitalize(),
                budget_str,
                prj.get("target_end_date") or "-",
            ]
            for idx, text in enumerate(vals):
                cell = row.cells[idx]
                set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(str(text))
                r.font.size = Pt(8.5)
